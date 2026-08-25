"""truth.document: format dispatch for CV uploads (PDF, DOCX, TXT, MD)."""

from __future__ import annotations

import io

import pytest

from truth.document import extract_document_text
from truth.pdf import DocumentExtractError


def test_txt_returns_decoded_text():
    text = extract_document_text("cv.txt", "Jane Doe\nSoftware Engineer".encode("utf-8"))
    assert "Jane Doe" in text
    assert "Software Engineer" in text


def test_md_returns_decoded_text():
    text = extract_document_text("cv.md", "# Jane Doe\n\n- Engineer".encode("utf-8"))
    assert "Jane Doe" in text
    assert "Engineer" in text


def test_docx_round_trips_paragraphs():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior Engineer at Acme")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_document_text("cv.docx", buf.getvalue())
    assert "Jane Doe" in text
    assert "Senior Engineer at Acme" in text


def test_unsupported_extension_raises_naming_supported_formats():
    with pytest.raises(DocumentExtractError) as exc:
        extract_document_text("cv.rtf", b"some bytes")
    message = str(exc.value)
    for fmt in ("PDF", "DOCX", "TXT", "MD"):
        assert fmt in message


def test_empty_bytes_raises():
    with pytest.raises(DocumentExtractError):
        extract_document_text("cv.pdf", b"")


def test_pdf_still_routes_to_pypdf_and_raises_on_garbage():
    with pytest.raises(DocumentExtractError):
        extract_document_text("cv.pdf", b"this is not a pdf at all")

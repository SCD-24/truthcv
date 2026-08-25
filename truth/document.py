"""Format dispatch for CV uploads: PDF, DOCX, TXT, and MD text extraction."""

from __future__ import annotations

from .pdf import DocumentExtractError, extract_text as _extract_pdf_text

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


def extension_for(filename: str) -> str:
    """Return the lowercased file extension (including the dot) of `filename`."""
    name = filename or ""
    idx = name.rfind(".")
    if idx == -1:
        return ""
    return name[idx:].lower()


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract paragraph and table cell text from a DOCX file."""
    import io

    try:
        from docx import Document
    except Exception as e:  # noqa: BLE001 — surfaced uniformly below
        raise DocumentExtractError(f"Could not read DOCX: {e}") from e
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:  # noqa: BLE001 — python-docx raises varied types
        raise DocumentExtractError(f"Could not read DOCX: {e}") from e
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_document_text(filename: str, file_bytes: bytes) -> str:
    """Extract text from an uploaded CV file, dispatching on its extension.

    Supports PDF (via truth.pdf), DOCX (via python-docx), and plain-text TXT/MD
    (decoded as UTF-8). Raises DocumentExtractError on empty input, an
    unsupported extension, an unreadable DOCX, or text that is empty once
    extracted.
    """
    if not file_bytes:
        raise DocumentExtractError("Uploaded file is empty.")
    ext = extension_for(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractError(
            "Unsupported file type. Upload your CV as PDF, DOCX, TXT, or MD."
        )
    if ext == ".pdf":
        text = _extract_pdf_text(file_bytes)
    elif ext == ".docx":
        text = _extract_docx_text(file_bytes)
    else:  # .txt / .md
        text = file_bytes.decode("utf-8", errors="replace")
    if not text.strip():
        raise DocumentExtractError(
            "No text found in the uploaded file. Upload your CV as PDF, DOCX, TXT, or MD."
        )
    return text
